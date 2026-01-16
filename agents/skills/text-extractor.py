#!/usr/bin/env python3
"""
Text Extractor Skill
Extract text from various document formats (PDF, DOCX, TXT, HTML)

Usage:
    from text_extractor import TextExtractor

    extractor = TextExtractor()
    result = extractor.execute("document.pdf", "pdf")
    print(result["text"])
"""

from typing import Dict, Any
from pathlib import Path
import logging

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class TextExtractor:
    """Extract text from various document formats"""

    def __init__(self):
        """Initialize text extractor"""
        self.supported_formats = ['pdf', 'docx', 'txt', 'html', 'md']

    def execute(self, file_path: str, format: str = None) -> Dict[str, Any]:
        """
        Extract text from document

        Args:
            file_path: Path to document
            format: Document format (pdf, docx, txt, html, md)
                   If not provided, inferred from file extension

        Returns:
            Dict with 'text' and 'metadata' keys:
            {
                "text": "extracted text content",
                "metadata": {
                    "format": "pdf",
                    "pages": 10,
                    "source": "path/to/file.pdf"
                },
                "success": True
            }

        Raises:
            ValueError: If format is not supported
            FileNotFoundError: If file doesn't exist
        """
        path = Path(file_path)

        # Validate file exists
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        # Infer format if not provided
        if format is None:
            format = path.suffix[1:].lower()  # Remove leading dot

        # Validate format
        if format not in self.supported_formats:
            raise ValueError(
                f"Unsupported format: {format}. "
                f"Supported formats: {', '.join(self.supported_formats)}"
            )

        logger.info(f"Extracting text from {file_path} (format: {format})")

        # Extract based on format
        extractor = getattr(self, f"_extract_{format}")
        result = extractor(path)

        # Add common metadata
        result["metadata"]["source"] = str(file_path)
        result["metadata"]["format"] = format
        result["success"] = True

        logger.info(f"Extracted {len(result['text'])} characters")

        return result

    def _extract_pdf(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from PDF"""
        try:
            import PyPDF2

            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                text = ""

                for page in reader.pages:
                    text += page.extract_text() + "\n"

                metadata = {
                    "pages": len(reader.pages),
                }

                # Add PDF metadata if available
                if reader.metadata:
                    metadata.update({
                        "author": reader.metadata.get('/Author', ''),
                        "title": reader.metadata.get('/Title', ''),
                        "created_date": reader.metadata.get('/CreationDate', '')
                    })

                return {
                    "text": text,
                    "metadata": metadata
                }

        except ImportError:
            raise ImportError(
                "PyPDF2 is required for PDF extraction. "
                "Install it with: pip install PyPDF2"
            )

    def _extract_docx(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from DOCX"""
        try:
            from docx import Document

            doc = Document(file_path)
            text = "\n".join([para.text for para in doc.paragraphs])

            metadata = {
                "paragraphs": len(doc.paragraphs),
            }

            # Add document properties if available
            core_props = doc.core_properties
            if core_props:
                metadata.update({
                    "author": core_props.author or '',
                    "title": core_props.title or '',
                    "created": str(core_props.created) if core_props.created else ''
                })

            return {
                "text": text,
                "metadata": metadata
            }

        except ImportError:
            raise ImportError(
                "python-docx is required for DOCX extraction. "
                "Install it with: pip install python-docx"
            )

    def _extract_txt(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from TXT file"""
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read()

        return {
            "text": text,
            "metadata": {
                "lines": text.count('\n') + 1,
                "characters": len(text)
            }
        }

    def _extract_html(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from HTML"""
        try:
            from bs4 import BeautifulSoup

            with open(file_path, 'r', encoding='utf-8') as file:
                soup = BeautifulSoup(file.read(), 'html.parser')

                # Remove script and style elements
                for element in soup(['script', 'style']):
                    element.decompose()

                text = soup.get_text(separator='\n', strip=True)

                # Get title
                title_tag = soup.find('title')
                title = title_tag.string if title_tag else ''

                return {
                    "text": text,
                    "metadata": {
                        "title": title,
                        "lines": text.count('\n') + 1
                    }
                }

        except ImportError:
            raise ImportError(
                "beautifulsoup4 is required for HTML extraction. "
                "Install it with: pip install beautifulsoup4"
            )

    def _extract_md(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from Markdown"""
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read()

        # Extract title from first # heading if present
        title = ''
        for line in text.split('\n'):
            if line.startswith('# '):
                title = line[2:].strip()
                break

        return {
            "text": text,
            "metadata": {
                "title": title,
                "lines": text.count('\n') + 1,
                "characters": len(text)
            }
        }


# Example usage
if __name__ == "__main__":
    import sys

    if len(sys.argv) < 2:
        print("Usage: python text-extractor.py <file_path> [format]")
        print(f"Supported formats: pdf, docx, txt, html, md")
        sys.exit(1)

    file_path = sys.argv[1]
    format = sys.argv[2] if len(sys.argv) > 2 else None

    extractor = TextExtractor()

    try:
        result = extractor.execute(file_path, format)
        print(f"Successfully extracted text from {file_path}")
        print(f"\nMetadata: {result['metadata']}")
        print(f"\nText preview (first 500 chars):")
        print(result['text'][:500])
        print("...")
    except Exception as e:
        logger.error(f"Error: {e}")
        sys.exit(1)
